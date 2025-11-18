"""
Serviço para processamento de arquivos.
"""

import os
import time
import mimetypes
from typing import Dict, Any, Optional, List
from werkzeug.utils import secure_filename
from config import config


class FileService:
    """Serviço para processamento de arquivos."""
    
    def __init__(self):
        """Inicializa o serviço de arquivos."""
        self.upload_folder = config.UPLOAD_FOLDER
        self.allowed_extensions = config.ALLOWED_EXTENSIONS
        self.max_content_length = config.MAX_CONTENT_LENGTH
        
        # Criar diretório de upload se não existir
        if not os.path.exists(self.upload_folder):
            os.makedirs(self.upload_folder)
    
    def is_allowed_file(self, filename: str) -> bool:
        """
        Verifica se o arquivo tem extensão permitida.
        
        Args:
            filename: Nome do arquivo
            
        Returns:
            True se o arquivo é permitido
        """
        if not filename:
            return False
        
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in self.allowed_extensions
    
    def get_file_extension(self, filename: str) -> str:
        """
        Obtém a extensão do arquivo.
        
        Args:
            filename: Nome do arquivo
            
        Returns:
            Extensão do arquivo (sem o ponto)
        """
        if not filename or '.' not in filename:
            return ''
        return filename.rsplit('.', 1)[1].lower()
    
    def get_file_mime_type(self, filename: str) -> str:
        """
        Obtém o tipo MIME do arquivo.
        
        Args:
            filename: Nome do arquivo
            
        Returns:
            Tipo MIME do arquivo
        """
        mime_type, _ = mimetypes.guess_type(filename)
        return mime_type or 'application/octet-stream'
    
    def save_file(self, file, filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Salva um arquivo no diretório de upload.
        
        Args:
            file: Objeto de arquivo do Flask
            filename: Nome personalizado (opcional)
            
        Returns:
            Dicionário com resultado da operação
        """
        try:
            if not file or not file.filename:
                return {
                    "success": False,
                    "error": "Nenhum arquivo fornecido"
                }
            
            if not self.is_allowed_file(file.filename):
                return {
                    "success": False,
                    "error": f"Tipo de arquivo não permitido. Extensões aceitas: {', '.join(self.allowed_extensions)}"
                }
            
            # Usar nome personalizado ou nome seguro do arquivo original
            if filename:
                safe_filename = secure_filename(filename)
            else:
                safe_filename = secure_filename(file.filename)
            
            # Garantir que o nome seja único
            file_path = os.path.join(self.upload_folder, safe_filename)
            counter = 1
            while os.path.exists(file_path):
                name, ext = os.path.splitext(safe_filename)
                file_path = os.path.join(self.upload_folder, f"{name}_{counter}{ext}")
                counter += 1
            
            # Salvar arquivo
            file.save(file_path)
            
            return {
                "success": True,
                "file_path": file_path,
                "filename": os.path.basename(file_path),
                "size": os.path.getsize(file_path),
                "mime_type": self.get_file_mime_type(file_path)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Erro ao salvar arquivo: {str(e)}"
            }
    
    def read_file_content(self, file_path: str) -> Dict[str, Any]:
        """
        Lê o conteúdo de um arquivo.
        
        Args:
            file_path: Caminho para o arquivo
            
        Returns:
            Dicionário com o conteúdo do arquivo
        """
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": "Arquivo não encontrado"
                }
            
            # Detectar tipo de arquivo e ler adequadamente
            mime_type = self.get_file_mime_type(file_path)
            
            if mime_type.startswith('text/'):
                # Arquivo de texto
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            else:
                # Arquivo binário - ler como bytes
                with open(file_path, 'rb') as f:
                    content = f.read()
            
            return {
                "success": True,
                "content": content,
                "mime_type": mime_type,
                "size": os.path.getsize(file_path)
            }
            
        except UnicodeDecodeError:
            return {
                "success": False,
                "error": "Erro de codificação: arquivo não é texto válido"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Erro ao ler arquivo: {str(e)}"
            }
    
    def delete_file(self, file_path: str) -> Dict[str, Any]:
        """
        Remove um arquivo.
        
        Args:
            file_path: Caminho para o arquivo
            
        Returns:
            Dicionário com resultado da operação
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return {
                    "success": True,
                    "message": "Arquivo removido com sucesso"
                }
            else:
                return {
                    "success": False,
                    "error": "Arquivo não encontrado"
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"Erro ao remover arquivo: {str(e)}"
            }
    
    def get_uploaded_files(self) -> List[Dict[str, Any]]:
        """
        Lista todos os arquivos no diretório de upload.
        
        Returns:
            Lista de informações sobre os arquivos
        """
        try:
            files = []
            for filename in os.listdir(self.upload_folder):
                file_path = os.path.join(self.upload_folder, filename)
                if os.path.isfile(file_path):
                    files.append({
                        "filename": filename,
                        "file_path": file_path,
                        "size": os.path.getsize(file_path),
                        "mime_type": self.get_file_mime_type(file_path),
                        "extension": self.get_file_extension(filename)
                    })
            return files
        except Exception as e:
            return []
    
    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extrai texto de um arquivo de forma inteligente.
        
        Args:
            file_path: Caminho para o arquivo
            
        Returns:
            Dicionário com o texto extraído
        """
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": "Arquivo não encontrado"
                }
            
            file_extension = self.get_file_extension(file_path)
            mime_type = self.get_file_mime_type(file_path)
            
            # Arquivos de texto simples
            if file_extension in ['txt', 'md'] or mime_type.startswith('text/'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return {
                    "success": True,
                    "text": content,
                    "method": "text_file"
                }
            
            # PDF - requer biblioteca adicional
            elif file_extension == 'pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                    return {
                        "success": True,
                        "text": text.strip(),
                        "method": "pdf_extraction"
                    }
                except ImportError:
                    return {
                        "success": False,
                        "error": "Biblioteca PyPDF2 não instalada para processar PDFs"
                    }
            
            # DOCX - requer biblioteca adicional
            elif file_extension == 'docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return {
                        "success": True,
                        "text": text.strip(),
                        "method": "docx_extraction"
                    }
                except ImportError:
                    return {
                        "success": False,
                        "error": "Biblioteca python-docx não instalada para processar DOCX"
                    }
            
            # Áudio - requer Whisper para transcrição
            elif file_extension in ['mp3', 'wav']:
                try:
                    from services.transcription_service import TranscriptionService
                    transcription_service = TranscriptionService()
                    result = transcription_service.transcribe_audio(file_path)
                    
                    if result.get('success'):
                        return {
                            "success": True,
                            "text": result.get('text', ''),
                            "method": "audio_transcription",
                            "transcription_info": {
                                "language": result.get('language', 'pt'),
                                "segments": result.get('segments', 0)
                            }
                        }
                    else:
                        return {
                            "success": False,
                            "error": result.get('error', 'Erro desconhecido na transcrição'),
                            "method": "audio_transcription"
                        }
                except ImportError:
                    return {
                        "success": False,
                        "error": "Biblioteca openai-whisper não está instalada. Execute: pip install openai-whisper torch"
                    }
                except Exception as e:
                    error_msg = str(e)
                    # Melhorar mensagens de erro específicas
                    if "ffmpeg" in error_msg.lower() or "ffprobe" in error_msg.lower():
                        return {
                            "success": False,
                            "error": "FFmpeg não encontrado no sistema. O Whisper requer FFmpeg instalado. Instale FFmpeg e adicione ao PATH do sistema. No Windows: winget install FFmpeg"
                        }
                    return {
                        "success": False,
                        "error": f"Erro ao transcrever áudio: {error_msg}"
                    }
            
            # DOC - requer biblioteca adicional
            elif file_extension == 'doc':
                try:
                    import subprocess
                    import tempfile
                    
                    # Tentar converter DOC para TXT usando antiword (Linux/Mac) ou catdoc
                    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp_file:
                        tmp_path = tmp_file.name
                    
                    try:
                        subprocess.run(['antiword', file_path], stdout=open(tmp_path, 'w'), check=True)
                        with open(tmp_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                        os.unlink(tmp_path)
                        return {
                            "success": True,
                            "text": content.strip(),
                            "method": "doc_extraction_antiword"
                        }
                    except (subprocess.CalledProcessError, FileNotFoundError):
                        try:
                            subprocess.run(['catdoc', file_path], stdout=open(tmp_path, 'w'), check=True)
                            with open(tmp_path, 'r', encoding='utf-8') as f:
                                content = f.read()
                            os.unlink(tmp_path)
                            return {
                                "success": True,
                                "text": content.strip(),
                                "method": "doc_extraction_catdoc"
                            }
                        except (subprocess.CalledProcessError, FileNotFoundError):
                            return {
                                "success": False,
                                "error": "Ferramentas antiword ou catdoc não encontradas para processar DOC"
                            }
                except Exception as e:
                    return {
                        "success": False,
                        "error": f"Erro ao processar DOC: {str(e)}"
                    }
            
            # Fallback: tentar ler como texto
            else:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    return {
                        "success": True,
                        "text": content,
                        "method": "fallback_text"
                    }
                except Exception as e:
                    return {
                        "success": False,
                        "error": f"Tipo de arquivo não suportado: {file_extension}"
                    }
                    
        except Exception as e:
            return {
                "success": False,
                "error": f"Erro ao extrair texto: {str(e)}"
            }
    
    def create_document(self, content: str, format_type: str, filename: str = None, user_stories: str = None, summary: str = None, document_title: str = None) -> Dict[str, Any]:
        """
        Cria um documento no formato especificado.
        
        Args:
            content: Conteúdo completo do documento (ou será montado de user_stories e summary)
            format_type: Tipo do documento ('pdf' ou 'docx')
            filename: Nome do arquivo (opcional)
            user_stories: Histórias de usuário (opcional, se fornecido será usado em vez de content)
            summary: Resumo da reunião (opcional)
            document_title: Título do documento (opcional, será determinado automaticamente se não fornecido)
            
        Returns:
            Dicionário com resultado da criação
        """
        try:
            # Determinar título do documento baseado no conteúdo
            if not document_title:
                if user_stories and summary:
                    # Tentar extrair título da primeira HU
                    title_from_hus = self._extract_title_from_hus(user_stories)
                    if title_from_hus:
                        document_title = f"{title_from_hus} - Histórias de Usuário e Resumo"
                    else:
                        document_title = "Histórias de Usuário e Resumo da Reunião"
                elif user_stories:
                    # Tentar extrair título da primeira HU
                    title_from_hus = self._extract_title_from_hus(user_stories)
                    if title_from_hus:
                        document_title = title_from_hus
                    else:
                        document_title = "Histórias de Usuário"
                elif summary:
                    document_title = "Resumo da Reunião"
                else:
                    document_title = "Documento Gerado"
            
            # Determinar nome do arquivo baseado no título
            if not filename:
                timestamp = int(time.time())
                # Criar nome de arquivo baseado no título
                safe_title = "".join(c for c in document_title if c.isalnum() or c in (' ', '-', '_')).strip()
                safe_title = safe_title.replace(' ', '_')[:50]  # Limitar tamanho
                if not safe_title:
                    safe_title = "documento"
                filename = f"{safe_title}_{timestamp}"
            
            # Montar conteúdo se user_stories e summary foram fornecidos
            if user_stories or summary:
                content_parts = []
                if user_stories:
                    content_parts.append(user_stories)
                if summary:
                    if content_parts:
                        content_parts.append("\n\n---\n\n")
                    content_parts.append(f"# Resumo da Reunião\n\n{summary}")
                content = "\n".join(content_parts)
            
            file_path = os.path.join(self.upload_folder, f"{filename}.{format_type}")
            
            if format_type.lower() == 'pdf':
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.units import inch
                    from reportlab.lib import colors
                    import re
                    
                    doc = SimpleDocTemplate(file_path, pagesize=letter, 
                                           rightMargin=72, leftMargin=72,
                                           topMargin=72, bottomMargin=72)
                    styles = getSampleStyleSheet()
                    story = []
                    
                    # Estilos customizados
                    title_style = ParagraphStyle(
                        'CustomTitle',
                        parent=styles['Title'],
                        fontSize=24,
                        textColor=colors.HexColor('#FF6F00'),
                        spaceAfter=30,
                        alignment=1  # Center
                    )
                    
                    heading_style = ParagraphStyle(
                        'CustomHeading',
                        parent=styles['Heading1'],
                        fontSize=18,
                        textColor=colors.HexColor('#2D2D2D'),
                        spaceAfter=12,
                        spaceBefore=20
                    )
                    
                    # Título principal (usar document_title)
                    title = Paragraph(document_title, title_style)
                    story.append(title)
                    story.append(Spacer(1, 20))
                    
                    # Processar conteúdo com markdown básico
                    lines = content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            story.append(Spacer(1, 6))
                            continue
                        
                        # Títulos
                        if line.startswith('### '):
                            text = line[4:].strip()
                            p = Paragraph(f"<b>{text}</b>", styles['Heading3'])
                            story.append(p)
                            story.append(Spacer(1, 6))
                        elif line.startswith('## '):
                            text = line[3:].strip()
                            p = Paragraph(text, heading_style)
                            story.append(p)
                            story.append(Spacer(1, 12))
                        elif line.startswith('# '):
                            text = line[2:].strip()
                            p = Paragraph(text, heading_style)
                            story.append(p)
                            story.append(Spacer(1, 12))
                        else:
                            # Converter markdown básico para HTML
                            text = line
                            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
                            text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
                            p = Paragraph(text, styles['Normal'])
                            story.append(p)
                            story.append(Spacer(1, 6))
                    
                    doc.build(story)
                    
                    return {
                        "success": True,
                        "file_path": file_path,
                        "filename": f"{filename}.pdf",
                        "size": os.path.getsize(file_path)
                    }
                    
                except ImportError:
                    return {
                        "success": False,
                        "error": "Biblioteca reportlab não instalada para criar PDFs"
                    }
            
            elif format_type.lower() in ['docx', 'doc']:
                try:
                    from docx import Document
                    from docx.shared import Inches, Pt, RGBColor
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    import re
                    
                    doc = Document()
                    
                    # Título principal (usar document_title)
                    title = doc.add_heading(document_title, 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title.runs[0]
                    title_run.font.color.rgb = RGBColor(255, 111, 0)
                    
                    # Processar conteúdo
                    lines = content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            doc.add_paragraph()
                            continue
                        
                        # Títulos
                        if line.startswith('### '):
                            text = line[4:].strip()
                            p = doc.add_heading(text, level=3)
                        elif line.startswith('## '):
                            text = line[3:].strip()
                            p = doc.add_heading(text, level=2)
                        elif line.startswith('# '):
                            text = line[2:].strip()
                            p = doc.add_heading(text, level=1)
                        else:
                            # Processar markdown básico
                            para = doc.add_paragraph()
                            # Dividir por formatação markdown
                            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                            for part in parts:
                                if part.startswith('**') and part.endswith('**'):
                                    run = para.add_run(part[2:-2])
                                    run.bold = True
                                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                                    run = para.add_run(part[1:-1])
                                    run.italic = True
                                elif part:
                                    para.add_run(part)
                    
                    # Salvar como .docx (mesmo se format_type for 'doc')
                    file_path_docx = file_path.replace('.doc', '.docx') if format_type.lower() == 'doc' else file_path
                    doc.save(file_path_docx)
                    
                    return {
                        "success": True,
                        "file_path": file_path_docx,
                        "filename": f"{filename}.docx",
                        "size": os.path.getsize(file_path_docx)
                    }
                    
                except ImportError:
                    return {
                        "success": False,
                        "error": "Biblioteca python-docx não instalada para criar DOCX"
                    }
            
            else:
                return {
                    "success": False,
                    "error": f"Formato não suportado: {format_type}. Use 'pdf' ou 'doc'"
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": f"Erro ao criar documento: {str(e)}"
            }
    
    def _extract_title_from_hus(self, user_stories: str) -> Optional[str]:
        """
        Extrai o título da primeira História de Usuário.
        
        Args:
            user_stories: Texto das Histórias de Usuário
            
        Returns:
            Título extraído ou None se não encontrado
        """
        try:
            lines = user_stories.split('\n')
            for i, line in enumerate(lines):
                line = line.strip()
                # Procurar por padrão "## 1. **Nome da História de Usuário**" seguido do nome
                if line.startswith('## 1.') and 'Nome da História de Usuário' in line:
                    # Próxima linha não vazia geralmente contém o nome
                    for j in range(i + 1, min(i + 5, len(lines))):
                        next_line = lines[j].strip()
                        if next_line and not next_line.startswith('#'):
                            # Remover formatação markdown
                            title = next_line.replace('**', '').replace('*', '').strip()
                            if title and len(title) > 5:  # Título válido
                                return title
                # Alternativa: procurar por linha que começa com "# História de Usuário"
                elif line.startswith('# ') and 'História de Usuário' in line:
                    # Extrair o nome após "História de Usuário"
                    parts = line.split(':', 1)
                    if len(parts) > 1:
                        title = parts[1].strip()
                        if title:
                            return title
                    # Ou usar a linha completa sem o "# "
                    title = line[2:].strip()
                    if title and len(title) > 5:
                        return title
            return None
        except Exception as e:
            print(f"[DEBUG] Erro ao extrair título das HUs: {str(e)}")
            return None
